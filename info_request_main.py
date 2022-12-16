"""Made By Alex Hyman 26/08/2021"""

# Imports 

import re # regex is required for finding patters within strings
from openpyxl import load_workbook # for opening excel workbooks
from zipfile import ZipFile # for opening zip files
import os # for accessing directories on the os 
from docx import Document 
from docx.shared import Cm # for creating word documents
from pdf2image import convert_from_path # for converting pdfs into images
from tkinter import * # for creating a graphical user interface
import tkinter.messagebox
from PIL import Image # for obtaining image dimensions
import shutil # for removing 'temp' directory after documents have been made
from pathlib import Path

# Constants

LETTERS_AND_PHOTOS_PATH = "Gordon and Noble\Gordon and Noble - Documents\Parking\Euro Car Parks\Euro Car Parks - Letters and Photos"

# Functions

def create_gui():
    # Create the root window
    root = Tk()
    root.title("Information Requests")
    root.geometry('600x600')

    # Create a listbox
    listbox = Listbox(root, width=400, height=500, selectmode=SINGLE)
    files = os.listdir("PCN_list")
    for i, file in enumerate(files, start=1):
        listbox.insert(i, file)

    tkinter.messagebox.showwarning('Warning', "Programme will not work if the first row of the spreadsheet has headings/titles. If the top row has titles please delete that row from the spreadsheet")
    tkinter.messagebox.showwarning('Warning', "Programme creates evidence based on the order it is given in spreadsheet. Sort the spreadsheet by column H if you want it in chronological order (typically this has already been done)")


    def selected_item():
        while len(listbox.curselection()) != 1: 
            print("nothing selected...")
            print("Terminating programme...")
            root.destroy()
            return None
        i = listbox.curselection()[0] 
        main(listbox.get(i), root)


    btn = Button(root, text='Create Evidence', command=selected_item)

    # Placing the button and listbox
    btn.pack(side='bottom')
    listbox.pack()

    root.mainloop()

def get_ids(sheet):
    # From an excel sheet, this returns a list of all the ids given in the second column
    ids, row_number, COL = [], 1, 2    
    while sheet.cell(row=row_number, column=COL).value != None: 
        ids.append(sheet.cell(row=row_number, column=COL).value)
        row_number += 1
    return ids

def get_dates(sheet):
    # from an excel sheet, this returns a list of all the dates given in the first column 
    # it formats them into a string of ddmmyyyy so we can concatenate with filenames
    dates, row_number, COL = [], 1, 1
    while sheet.cell(row=row_number, column=COL).value != None: 
        date = sheet.cell(row=row_number, column=COL).value
        if isinstance(date, str) and date[-1] == 'b':
            # This is an edge case where we are given 30/01/2020 b as a string and need 30012020B
            dates.append(date.replace("1", "01").replace("/","").replace(" b", "B"))
        else:
            dates.append(date.strftime(r'%d%m%Y'))
        row_number += 1
    return dates

def alter_dates_and_ids(ids, dates):
    # There seems to be a bug in some Spreadsheets where the date 01062020 is ambiguous 
    # And it could either be 03062020 or 10062020
    # This function replaces the erroneous date with its 2 contenders and aligns a parallel id array
    newDates, newIds = [], []
    for i, date in zip(ids, dates): 
        if date == "01062020":
            newDates.append("03062020") # replace erroneous date with 2 contenders and then duplicate in id array
            newDates.append("10062020")
            newIds.append(i)
            newIds.append(i)
        else:
            newDates.append(date)
            newIds.append(i)
    return newDates, newIds

def get_txt_file(dateString) -> str:
    # given a dateString, this returns a txt file as a string
    base = r"C:\Users" + os.sep + os.getlogin() + os.sep + LETTERS_AND_PHOTOS_PATH + os.sep + "AttachmentIndex_"
    path = base + dateString + ".txt"

    file = Path(path)
    
    if file.is_file():
        with open(path, 'r') as f: 
            return f.read()
    else:
        # The text file does not exist
        return None

def get_list_of_files_with_id(s: str, i: int) -> list:
    # s will be the string we are searching over for the string i (which is the id)
   
    regex = r"(^" + str(i) + ".*$)+" # this finds all lines with the id number on it
    
    matches = re.finditer(regex, s, re.MULTILINE) # save all lines which contain id into a matches variable
    
    # Concatenate all lines with the id into one string
    next_test = ""
    for match in matches:
        next_test += match.group(0) + "\n"
    
    # Now we want to remove most of the crap around the string, leaving just the filenames
    regex2 = r".+\^.+\^(.*)\^.+\n" 
    
    # Remove the gunk and seperate filenames by commas
    sub = re.sub(regex2, r"\1,", next_test)
    
    # Return a list of all the filenames
    # we dont return final element since it is an empty string 
    return sub.split(",")[:-1] 

def get_zip_path(dateString: str) -> str:
    # Given a date string, returns the corresponding zip folder path
    base = r"C:\Users" + os.sep + os.getlogin() + os.sep + LETTERS_AND_PHOTOS_PATH
    zip_name = r"\Attachments_" + dateString + ".zip"
    return base + zip_name

def file_in_zip(path, filename) -> bool:
    # Create a ZipFile Object
    with ZipFile(path, 'r') as zipObj:
    # Get list of files names in zip
        listOfiles = zipObj.namelist()
        # Iterate over the list of file names in given list & print them
        for elem in listOfiles:
            if elem == filename:
                return True
    return False

def all_files_in_zip(path, files) -> bool:
    if len(files) == 0: 
        return False
    
    for file in files: 
        if not file_in_zip(path, file):
            return False
    return True

def extract_files(zip_path, files, temp_folder_name = 'temp') -> None:
    # given the path to a zip folder, extract all the files given
    flag = True
    with ZipFile(zip_path, 'r') as zipObject:
        listOfFileNames = zipObject.namelist()
        for filename in listOfFileNames:
            if filename_in_files(filename, files):
                # Extract a single file from zip
                zipObject.extract(filename, temp_folder_name)
                print('file extracted')

def convert_pdf(path, filename):
    images = convert_from_path(path)
    for i in range(len(images)):
        images[i].save(filename + str(i) + ".jpg", "JPEG")

def filename_in_files(filename, files):
    # check whether a given filename contains filenames from the list files
    for file in files:
        if file in filename:
            return True
    return False

def get_workbook_name() -> str:
    files = os.listdir("PCN_list")
    for file in files:
        print(file)
    return "PCN_list" + os.sep + input("Please enter the name of the workbook you like to use in the PCN List folder: ")

def create_narrow_document() -> Document:
    document = Document() 

    #changing the page margins
    sections = document.sections
    margin = 1.27
    for section in sections:
        section.top_margin = Cm(margin)
        section.bottom_margin = Cm(margin)
        section.left_margin = Cm(margin)
        section.right_margin = Cm(margin)
    return document

def convert_and_insert_pdfs(document, folderPath):
    n=1
    for r, d, f in os.walk(folderPath):
        for item in f: 
            if '.pdf' in item: 
                convert_pdf(folderPath + "\\" + item, str(n) + "_")
                document.add_picture(str(n) + "_0.jpg", height=Cm(26.08), width=Cm(18.44)) # there is only one page so the number will be 0
                n += 1

def insert_images(document, folderPath):
    number_of_files = number_of_files_in_directory(folderPath)
    if number_of_files > 6: 
        H, W = 6, 8
    else:
        H, W = 9.8, 13
    for r, d, f in os.walk(folderPath):
        for item in f: 
            if '.jpg' in item: 
                img = Image.open(r + os.sep + item)
                width, height = img.size 
                area = width*height

                if area > 30_000: 
                    document.add_picture(r + os.sep + item, height=Cm(H), width=Cm(W))
                else:
                    document.add_picture(r + os.sep + item, height=Cm(1.3), width=Cm(3.89))

def number_of_files_in_directory(path: str) -> int:
    n = 0 
    for _ in os.walk(path):
        n += 1
    return n

def create_evidence_documents(filename: str) -> None:
    i = 0
    n = number_of_files_in_directory('temp')
    os.mkdir(filename)
    while i <= n//10: # work into batches of ten, i keeps track of the number of batches of ten
        document = create_narrow_document()
        for j in range(min(10, n - i*10)): # we want to create documents of 10 or less pieces of evidence
            k = 10*i + j + 1 # This is which number of the piece of evidence we are using
            path = 'temp' + os.sep + str(k)
            convert_and_insert_pdfs(document, path)
            insert_images(document, path)
        document.save(filename + os.sep + filename + " evidence pt " + str(i+1) + '.docx')            
        i += 1

def delete_temporary_files_and_directories():
    print("deleting temporary files...")
    try:
        shutil.rmtree("temp")
        os.remove("1_0.jpg")
        os.remove("2_0.jpg")
        print("Successfully deleted temp folder and jpegs")
    except OSError as e:
        print(f"Error: {e.strerror}")

def main(workbook_name, gui):
    
    try: 
        wb = load_workbook("PCN_list" + os.sep + workbook_name, read_only=True)
    except FileNotFoundError as e: 
        print("File not found in directory...")
        print("Exiting programme")
        return

    sheet = wb.active # The workbooks should only have one sheet 

    ids = get_ids(sheet) # parallel arrays of ids and their corresponding dates
    dates = get_dates(sheet)
    dates, ids = alter_dates_and_ids(ids, dates)
    temp_folder_name = 1

    found_ids = set() 

    for idNumber, dateString in zip(ids, dates):
        
        print(f"id: {idNumber}, date: {dateString}")
        text_file = get_txt_file(dateString) # get the text file associated with the dateString
        
        if text_file is None: 
            print("Text file was not found\n")
            continue
        
        files = get_list_of_files_with_id(text_file, idNumber) # get all filenames which include the id
        zip_path = get_zip_path(dateString) # get the zip folder associated with the date string

        # extract all files from zip if they exist
        if all_files_in_zip(zip_path, files):
            extract_files(zip_path, files, 'temp' + os.sep + str(temp_folder_name))
            temp_folder_name += 1

            found_ids.add(idNumber)
        else:
            print(f"The following id: {idNumber}, is missing at this date: {dateString}")
            
        print("")

    missing_ids = {i for i in ids if i not in found_ids}
    if len(missing_ids) == 0: 
        print("All ids were found!")
    else:
        print("The following ids were unable to be found")
        for i in missing_ids:
            print(i)
    print()

    name = str(sheet.cell(2, 11).value)
    reg  = str(sheet.cell(2, 13).value) 
    
    create_evidence_documents(name + " " + reg) # create the evidence documents from images in temp folder
    delete_temporary_files_and_directories()
    print("DONE")
    tkinter.messagebox.showinfo("Process finished", "Finished making evidence files")
    gui.destroy()
    
if __name__ == "__main__": create_gui()